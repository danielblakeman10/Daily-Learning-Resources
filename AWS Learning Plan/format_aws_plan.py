from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Styles setup
styles = doc.styles

h1 = styles['Heading 1']
h1.font.name = 'Arial'
h1.font.size = Pt(22)
h1.font.color.rgb = RGBColor(0, 82, 204)
h1.font.bold = True

h2 = styles['Heading 2']
h2.font.name = 'Arial'
h2.font.size = Pt(16)
h2.font.color.rgb = RGBColor(0, 102, 255)
h2.font.bold = True

title = doc.add_heading('AWS DevOps Learning Path', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Beginner to Enterprise-Level', style='Subtitle')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Introduction
p = doc.add_paragraph()
r = p.add_run('Overview: 4 weeks to cover 5 Core Pillars')
r.bold = True
r.font.size = Pt(12)

pillars = [
    'Foundation & Governance',
    'Infrastructure & Networking [Common for AWS + AZ]',
    'Compute & Application Platforms',
    'DevOps & Automation (CORE FOCUS)',
    'Observability, Security & Reliability'
]
for pillar in pillars:
    doc.add_paragraph(pillar, style='List Bullet')

doc.add_paragraph()
p2 = doc.add_paragraph()
p2.add_run('Each AWS service should map into real DevOps workflows like:').bold = True
workflows = [
    'CI/CD pipelines',
    'Infrastructure as Code (IaC)',
    'Container orchestration',
    'Monitoring & incident response',
    'Security & compliance automation'
]
for w in workflows:
    doc.add_paragraph(w, style='List Bullet')

doc.add_paragraph()
goal = doc.add_paragraph()
goal_run = goal.add_run('Final Goal: "Design, build, automate, secure, and operate production-grade AWS systems using DevOps principles."')
goal_run.bold = True
goal_run.italic = True
goal_run.font.color.rgb = RGBColor(0, 153, 0)
goal.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

sections = [
    {
        "title": "1. Foundation & Multi-Account Governance (START HERE)",
        "goal": "Understand how enterprises structure AWS",
        "topics": ["AWS Global Infrastructure", "Multi-account strategy", "Landing Zone design", "Governance & guardrails"],
        "services": ["AWS Organizations", "Control Tower", "IAM Identity Center"],
        "devops": ["Separate accounts for Dev / QA / Prod", "Centralized security + logging"]
    },
    {
        "title": "2. Identity, Security & Compliance (CRITICAL FOR DEVOPS)",
        "goal": "Secure everything by design",
        "topics": ["IAM deep dive (RBAC, ABAC)", "Temporary credentials (STS)", "Secrets management", "Zero trust access"],
        "services": ["IAM", "STS", "Secrets Manager", "KMS"],
        "devops": ["CI/CD roles assume permissions", "No hardcoded credentials", "Secure pipeline execution"]
    },
    {
        "title": "3. Networking & Hybrid Architecture",
        "goal": "Build enterprise network architecture",
        "topics": ["VPC design (Hub-Spoke)", "Private vs Public workloads", "Hybrid connectivity"],
        "services": ["VPC", "Transit Gateway", "PrivateLink", "Route 53"],
        "devops": ["Private CI/CD runners", "Secure service-to-service communication"]
    },
    {
        "title": "4. Compute & Application Hosting Models",
        "goal": "Choose the right runtime",
        "topics": ["VM-based (EC2)", "Container-based (ECS/EKS)", "Serverless (Lambda)"],
        "services": ["EC2 + Auto Scaling", "ALB/NLB", "ECS / EKS", "Lambda"],
        "devops": ["Blue/Green deployments", "Auto-scaling apps", "Immutable infrastructure"]
    },
    {
        "title": "5. Storage & Data Strategy",
        "goal": "Understand persistence & performance",
        "topics": ["Object vs Block vs File", "Backup & disaster recovery", "Data lifecycle"],
        "services": ["S3", "EBS / EFS", "FSx"],
        "devops": ["Artifact storage (S3)", "Backup automation", "Log storage"]
    },
    {
        "title": "6. Databases & Data Layer",
        "goal": "Managed DB operations",
        "topics": ["HA design (Multi-AZ)", "Read scaling", "NoSQL vs SQL"],
        "services": ["RDS / Aurora", "DynamoDB", "ElastiCache"],
        "devops": ["Automated backups", "DB migration pipelines"]
    },
    {
        "title": "7. Serverless & Event-Driven Architecture",
        "goal": "Build loosely coupled systems",
        "topics": ["Event-driven design", "Async communication"],
        "services": ["Lambda", "API Gateway", "SQS / SNS", "EventBridge", "Step Functions"],
        "devops": ["Event-triggered pipelines", "Microservices communication"]
    },
    {
        "title": "8. DevOps & Automation (MOST IMPORTANT)",
        "goal": "Build and release software reliably",
        "topics": ["CI/CD pipelines", "Infrastructure as Code", "GitOps", "Release strategies"],
        "services": ["CodePipeline / CodeBuild / CodeDeploy", "Terraform (VERY IMPORTANT)", "CloudFormation", "Systems Manager"],
        "devops": ["Build -> Test -> Deploy pipeline", "Blue/Green & Canary deployments", "Automated rollback"]
    },
    {
        "title": "9. Containers & Kubernetes (MODERN DEVOPS CORE)",
        "goal": "Orchestrate microservices",
        "topics": ["Docker fundamentals", "Container orchestration", "Service mesh basics"],
        "services": ["ECR", "ECS / Fargate", "EKS"],
        "devops": ["Microservices deployment", "CI/CD to Kubernetes", "Helm + GitOps"]
    },
    {
        "title": "10. Observability & Monitoring",
        "goal": "Operate production systems",
        "topics": ["Metrics, logs, traces", "Alerting", "Incident response"],
        "services": ["CloudWatch", "CloudTrail", "AWS Config", "Managed Prometheus / Grafana"],
        "devops": ["SLA monitoring", "Root cause analysis", "Auto-remediation"]
    },
    {
        "title": "11. Security Operations (SecDevOps)",
        "goal": "Continuous security",
        "topics": ["Automated threat detection", "WAF and edge security", "Vulnerability management"],
        "services": ["GuardDuty", "Inspector", "Security Hub", "WAF / Shield"],
        "devops": ["DevSecOps pipelines", "Vulnerability scanning", "Compliance automation"]
    },
    {
        "title": "12. Migration, Cost & Optimization",
        "goal": "Enterprise efficiency",
        "topics": ["Migration strategies (6Rs)", "Cost optimization", "Performance tuning"],
        "services": ["DMS", "Migration Hub", "Cost Explorer", "Savings Plans"],
        "devops": ["FinOps", "Automated scaling to zero"]
    }
]

for section in sections:
    doc.add_heading(section["title"], level=2)
    
    goal_p = doc.add_paragraph()
    goal_p.add_run("Goal: ").bold = True
    gr = goal_p.add_run(section["goal"])
    gr.italic = True
    gr.font.color.rgb = RGBColor(80, 80, 80)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Shading Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Core Topics'
    hdr_cells[1].text = 'Key Services'
    hdr_cells[2].text = 'DevOps Mapping'
    hdr_cells[3].text = 'My Notes'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                
    max_len = max(len(section["topics"]), len(section["services"]), len(section["devops"]))
    
    for i in range(max_len):
        row_cells = table.add_row().cells
        if i < len(section["topics"]):
            row_cells[0].text = "• " + section["topics"][i]
        if i < len(section["services"]):
            row_cells[1].text = "• " + section["services"][i]
        if i < len(section["devops"]):
            row_cells[2].text = "• " + section["devops"][i]
        # Add a few blank lines so the cell is spacious for written notes
        row_cells[3].text = "\n\n\n"
            
    doc.add_paragraph()

doc.core_properties.author = "Daniel Blakeman"
doc.core_properties.last_modified_by = "Daniel Blakeman"

output_path = r"C:\Users\Daniel Blakeman\Documents\GitHub Repos\Learning Resources\AWS Learning Plan\New AWS Learning Plan.docx"
try:
    doc.save(output_path)
    print("Saved cleanly formatted doc to:", output_path)
except PermissionError:
    # If the user has it open, save to a new file so it doesn't crash
    alt_path = output_path.replace(".docx", " - Workbook.docx")
    doc.save(alt_path)
    print("Original file was open, saved as:", alt_path)
    import os
    os.system(f'start "" "{alt_path}"')
